from docx import Document
from docx.table import Table
from docx.shared import RGBColor
import xml.etree.ElementTree as ET
import re   
import zipfile
import xml.etree.ElementTree as ET

# -------------------------------
# Extract titles by color
# -------------------------------

def extract_titles_by_color_from_docx(docx_file):
    doc = Document(docx_file)
    titles = {"red_titles": [], "grey_titles": []}

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Title":
            for run in paragraph.runs:
                if run.font.color and run.font.color.rgb:
                    if run.font.color.rgb == RGBColor(0xD6, 0x1F, 0x26):
                        titles["red_titles"].append(run.text.strip())
                    elif run.font.color.rgb == RGBColor(0x92, 0x94, 0x95):
                        titles["grey_titles"].append(run.text.strip())

    return {
        "red_titles": " ".join(titles["red_titles"]),
        "grey_titles": " ".join(titles["grey_titles"])
    }


# -------------------------------
# Extract journal metadata
# -------------------------------
def extract_journal_meta(docx_file):
    doc = Document(docx_file)
    journal_meta_string = ""

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if (
                run.bold
                and not run.italic
                and not run.underline
                and run.font.name == "Comic Sans MS"
                and run.font.size == 88900
                and run.font.color.rgb == RGBColor(0x13, 0x16, 0x17)
            ):
                journal_meta_string = paragraph.text.strip()

    if not journal_meta_string:
        
        return None
    
    return {
        "journal_title": journal_meta_string.split("-")[0].strip(),
        "journal_year": journal_meta_string.split("-")[1].split(";")[0].strip(),
        "jounral_id": journal_meta_string.split("/")[1].split("-")[0].split(".")[0].strip(),
        "journal_volume": journal_meta_string.split("Vol")[1].split(" ")[0].strip(),
        "journal_issue": journal_meta_string.split("(")[1].split(")")[0].strip(),
        "journal_start_page": journal_meta_string.split(":")[1].split("-")[0].strip(),
        "journal_end_page": journal_meta_string.split(":")[1].split("-")[1].strip().split("\t")[0],
        "journal_article_id": journal_meta_string.split("DOI:")[1].strip()
    }

# extract abstract 
def extract_abstract(docx_file):
    with zipfile.ZipFile(docx_file) as docx_zip:
        xml_content = docx_zip.read("word/document.xml")
        root = ET.fromstring(xml_content)
        arr = []
        # Namespace for WordprocessingML text elements
        w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for drawing in root.findall(f".//{w_ns}drawing"):
            # find all text elements under the drawing (<w:t>) and append their text
            for t in drawing.findall(f".//{w_ns}t"):
                if t.text:
                    arr.append(t.text)
    fr_abstract_text = " ".join(arr[2:arr.index("Mots")]).strip()
    fr_keywords = [k.strip() for k in " ".join(arr[arr.index("Mots")+4:arr.index("BSTRACT")-1]).strip().split(",") if k.strip()]

    en_abstract_text = " ".join(arr[arr.index("BSTRACT")+1:arr.index("Key")]).strip()

    en_full_keywords_array = arr[arr.index("Key")+4:]
    if "Figure" in en_full_keywords_array:
        if en_full_keywords_array.index("Figure") == -1:
            en_keywords = en_full_keywords_array
        else:
            en_keywords = en_full_keywords_array[:en_full_keywords_array.index("Figure")]
    else:
        en_keywords = en_full_keywords_array
    en_keywords = [k.strip() for k in " ".join(en_keywords).strip().split(",") if k.strip()]
    return     {
        "fr": {
            "title": "RESUME",
            "para": fr_abstract_text,
            "keywords": fr_keywords
        },
        "en": {
            "title": "ABSTRACT",
            "para": en_abstract_text,
            "keywords": en_keywords
        }
    }


#extract permissions info
def extract_permissions_info(docx_file):
    doc = Document(docx_file)
    permissions_info = {}
 
    permissions_info["copyright_year"] = 2022  # Assuming 2022 for this example
    permissions_info["license_type"] = 'open-access'  # Assuming open-access for this example
    
    permissions_info["license_url"] = ' https://creativecommons.org/licenses/by-nc-nd/4.0/'  # Assuming this URL for this example
    

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if (run.font.size == 76200
                and run.font.color.rgb == RGBColor(0x13, 0x16, 0x17)
            ):
                permissions_info["license_text"] = paragraph.text.strip()
    return permissions_info




# -------------------------------
# Extract contributors
# -------------------------------
def extract_contributors(docx_file):
    doc = Document(docx_file)
    contributors_list = []

    contributors = doc.paragraphs[5].text.split(",")

    for contributor in contributors:
        contributor = contributor.strip()
        if contributor[-1].isdigit():
            ref = contributor[-1]
            name_parts = contributor[:-1].split()
        else:
            ref = "None"
            name_parts = contributor.split()

        contributors_list.append({
            "name": name_parts[0],
            "surname": " ".join(name_parts[1:]),
            "ref": ref
        })

    return contributors_list
#extract affiliations before email
def extract_affiliations_before_email(docx_file):
    doc = Document(docx_file)
    res = {}
    var = "Correspondance"
    aff_id = -2
    for i, paragraph in enumerate(doc.paragraphs):
        if var in paragraph.text:
            aff_id += 1
            res[aff_id] = paragraph.text
            for paragraphe in doc.paragraphs[i+1:]:
                if "Email" in paragraphe.text:
                    break
                else:
                    aff_id += 1
                    res[aff_id] = paragraphe.text
            # Return dict skipping first 2 entries
            return {k: v for k, v in list(res.items())[2:]}
         
                                                        
#extract keys

#extract body
def extract_body(docx_file,):
    doc = Document(docx_file)
    sections = []
    current_sec = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            if paragraph.text.strip().upper() == "REFERENCE" or paragraph.text.strip().upper() == "RÉFÉRENCE":
                # Stop processing if heading is RÉFÉRENCE
                break
            if current_sec:
                sections.append(current_sec)
            current_sec = {
                "title": paragraph.text.strip(),
                "p": []
            }
        
        else:
            if current_sec is not None:
                if paragraph.text.strip():
                    text = paragraph.text.strip()
                    # Find all numbers inside parentheses
                    

                    parts = []
                    last = 0
                    for m in re.finditer(r'\((\d+)\)', text):
                        parts.append(text[last:m.start()])   # normal text
                        parts.append(("xref", m.group(1)))   # mark xref
                        last = m.end()

                    parts.append(text[last:])  # remaining text
                    current_sec["p"].append(parts)


    if current_sec:
        sections.append(current_sec)

    return sections
        

    



# -------------------------------
# Create XML
# -------------------------------
def create_xml_with_title(titles, journal, contributors, affiliations, permissions_info, abstract_data, sections, output_file):

    article = ET.Element("article", {
        "xml:lang": "fr",
        "xmlns:xlink": "http://www.w3.org/1999/xlink",
        "xmlns:mml": "http://www.w3.org/1998/Math/MathML",
        "article-type": "research-article"
    })

    front = ET.SubElement(article, "front")

    # journal-meta
    journal_meta = ET.SubElement(front, "journal-meta")
    ET.SubElement(journal_meta, "journal-id", {"journal-id-type": "publisher-id"}).text = journal["jounral_id"]

    journal_title_group = ET.SubElement(journal_meta, "journal-title-group")
    ET.SubElement(journal_title_group, "journal-title").text = journal["journal_title"]

    ET.SubElement(journal_meta, "issn", {"pub-type": "epub"}).text = "2724-7031"
    ET.SubElement(journal_meta, "issn", {"pub-type": "ppub"}).text = "0041-4131"

    # article-meta
    article_meta = ET.SubElement(front, "article-meta")
    ET.SubElement(article_meta, "article-id", {"pub-id-type": "doi"}).text = journal["journal_article_id"]

    title_group = ET.SubElement(article_meta, "title-group")
    ET.SubElement(title_group, "article-title").text = titles["grey_title"]

    trans_title_group = ET.SubElement(title_group, "trans-title-group")
    ET.SubElement(trans_title_group, "trans-title", {"xml:lang": "en"}).text = titles["red_title"]
    
    # contributors
    contrib_group = ET.SubElement(article_meta, "contrib-group")

    for contributor in contributors:
        contrib = ET.SubElement(contrib_group, "contrib", {"contrib-type": "author"})
        name = ET.SubElement(contrib, "name")
        ET.SubElement(name, "surname").text = contributor["surname"]
        ET.SubElement(name, "given-names").text = contributor["name"]

        if contributor["ref"] != "None":
            xref = ET.SubElement(contrib, "xref", {
                "ref-type": "aff",
                "rid": f"aff{contributor['ref']}"
            })
            xref.text = contributor["ref"]

    # affiliations

    for aff_id, aff_text in affiliations.items():
        aff = ET.SubElement(article_meta, "aff", {"id": f"aff{aff_id}"})
        ET.SubElement(aff, "institution").text = aff_text

    #pub
    pub=ET.SubElement(article_meta, "pub-date", {"pub-type": "ppub"})
    ET.SubElement(pub, "month").text = "12"
    ET.SubElement(pub, "year").text = journal["journal_year"]
    epub=ET.SubElement(article_meta, "pub-date", {"pub-type": "epub"})
    ET.SubElement(epub, "day").text = "05"
    ET.SubElement(epub, "month").text = "12"
    ET.SubElement(epub, "year").text = journal["journal_year"]
    ET.SubElement(article_meta, "volume").text = journal["journal_volume"]
    ET.SubElement(article_meta, "issue").text = journal["journal_issue"]
    ET.SubElement(article_meta, "fpage").text = journal["journal_start_page"]
    ET.SubElement(article_meta, "lpage").text = journal["journal_end_page"]
    permissions = ET.SubElement(article_meta, "permissions")
    copyright_year = ET.SubElement(permissions, "copyright-year").text = "2022"
    license = ET.SubElement(permissions, "license", {"license-type": "open-access", "xlink:href": "https://creativecommons.org/licenses/by-nc-nd/4.0/"})
    license_p = ET.SubElement(license, "license-p").text = permissions_info["license_text"]
    
    
    #  abstract
    abstract = ET.SubElement(article_meta, "abstract", {"abstract-type": "section"})
    title=ET.SubElement(abstract, "title").text = abstract_data["fr"]["title"]
    para=ET.SubElement(abstract, "p").text = abstract_data["fr"]["para"]
    trans=ET.SubElement(abstract, "trans-abstract", {"xml:lang": "en"})
    title_trans=ET.SubElement(trans, "title").text = abstract_data["en"]["title"]
    para_trans=ET.SubElement(trans, "p").text = abstract_data["en"]["para"]
    kwd=ET.SubElement(article_meta, "kwd-group", {"kwd-group-type": "author-keywords"})
    title_kwd=ET.SubElement(kwd, "title").text = "Mots clés"


    for key in abstract_data["fr"]["keywords"]:
        kwd_item=ET.SubElement(kwd, "kwd").text = key
    trans_kwd=ET.SubElement(article_meta, "kwd-group", {"kwd-group-type": "author-keywords", "xml:lang": "en"})
    title_trans_kwd=ET.SubElement(trans_kwd, "title").text = "Keywords"
    for key in abstract_data["en"]["keywords"]:
        kwd_item_trans=ET.SubElement(trans_kwd, "kwd").text = key   




    # body
    body = ET.SubElement(article, "body")
    table_counter = 0
    for section in sections:
        sec = ET.SubElement(body, "sec")
        ET.SubElement(sec, "title").text = section["title"]
        for para in section["p"]:
            p_elem = ET.SubElement(sec, "p")
            first = True

            for part in para:
                
                if isinstance(part, tuple) and part[0] == "xref":
                    num = part[1]
                    xref = ET.SubElement(p_elem, "xref", {
                        "ref-type": "bibr",
                        "rid": f"bib{num}"
                    })
                    xref.text = num
                    first = False
                
                else:
                    if first:
                        p_elem.text = part
                        first = False
                    else:
                        p_elem[-1].tail = part


    #back

    back = ET.SubElement(article, "back")
    ref_list = ET.SubElement(back, "ref-list")
    title_ref = ET.SubElement(ref_list, "title").text = "References"
    # Example reference
    ref = ET.SubElement(ref_list, "ref", {"id": "bib1"})
    element_citation = ET.SubElement(ref, "element-citation", {"publication-type": "book"})
    person_group = ET.SubElement(element_citation, "person-group", {"person-group-type": "author"})
    name = ET.SubElement(person_group, "name")
    ET.SubElement(name, "surname").text = "Eckel"
    ET.SubElement(name, "given-names").text = "RH"
    ET.SubElement(element_citation, "chapter-title").text = "The Metabolic Syndrome"
    ET.SubElement(element_citation, "source").text = "Harrison’s Principles of Internal Medicine"
    ET.SubElement(element_citation, "edition").text = "20e éd"
    ET.SubElement(element_citation, "publisher-name").text = "McGraw-Hill Education"
    ET.SubElement(element_citation, "publisher-loc").text = "New York, NY"
    ET.SubElement(element_citation, "year").text = "2018"
    ET.SubElement(element_citation, "comment").text = "[cité 26 août 2024]"
    ET.SubElement(element_citation, "ext-link", {"ext-link-type": "uri"}).text = " https://accessmedicine.mhmedical.com/content.aspx?aid=1183992857"    


    # write file
    tree = ET.ElementTree(article)
    with open(output_file, "wb") as f:
        f.write(b'<?xml version="1.0" encoding="UTF-8"?>\n')
        f.write(b'<!DOCTYPE article PUBLIC "-//NLM//DTD JATS (Z39.96) Journal Publishing DTD v1.1d1 20130915//EN" '
                b'"http://jats.nlm.nih.gov/publishing/1.1d1/JATS-journalpublishing1.dtd">\n')
        tree.write(f, encoding="utf-8", xml_declaration=False)


# -------------------------------
# Main
# -------------------------------
def main():
    input_docx = "article.docx"
    output_xml = "output.xml"

    titles_list = extract_titles_by_color_from_docx(input_docx)
    journal = extract_journal_meta(input_docx)
    contributors = extract_contributors(input_docx)
    affiliations = extract_affiliations_before_email(input_docx)
    permissions_info = extract_permissions_info(input_docx)
    abstract_data = extract_abstract(input_docx)
    paragraphs_body = extract_body(input_docx)
    titles = {
        "red_title": titles_list["red_titles"],
        "grey_title": titles_list["grey_titles"]
    }

    if not titles["red_title"] or not titles["grey_title"]:
        print("Titles not found.")
        return

    create_xml_with_title(titles, journal, contributors, affiliations, permissions_info, abstract_data, paragraphs_body, output_xml)
    print(f"XML file created successfully: {output_xml}")
    

    print("this is main branch:")
if __name__ == "__main__":
    main()
